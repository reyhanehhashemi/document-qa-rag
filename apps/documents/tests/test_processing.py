import tempfile
from io import BytesIO

from django.core.files.base import ContentFile
from django.test import TestCase, override_settings
from docx import Document as WordDocument

from apps.documents.models import Document
from apps.documents.services.exceptions import InvalidDocxError
from apps.documents.services.processing import process_document


class DocumentProcessingTests(TestCase):
    def setUp(self):
        self.temporary_media_directory = tempfile.TemporaryDirectory()

        self.media_override = override_settings(
            MEDIA_ROOT=self.temporary_media_directory.name,
        )

        self.media_override.enable()

    def tearDown(self):
        self.media_override.disable()
        self.temporary_media_directory.cleanup()

    def create_docx_bytes(self):
        output = BytesIO()

        word_document = WordDocument()

        word_document.add_paragraph(
            "Document title"
        )

        word_document.add_paragraph(
            "This document contains useful information."
        )

        word_document.save(output)

        return output.getvalue()

    def test_process_document_extracts_text(self):
        document = Document.objects.create(
            title="Test Document",
            file=ContentFile(
                self.create_docx_bytes(),
                name="test-document.docx",
            ),
        )

        process_document(document)

        document.refresh_from_db()

        self.assertEqual(
            document.status,
            Document.Status.PROCESSED,
        )

        self.assertIn(
            "Document title",
            document.text_content,
        )

        self.assertIn(
            "This document contains useful information.",
            document.text_content,
        )

        self.assertEqual(
            document.processing_error,
            "",
        )

    def test_invalid_docx_marks_document_as_failed(self):
        document = Document.objects.create(
            title="Invalid Document",
            file=ContentFile(
                b"This is not a real DOCX file.",
                name="invalid.docx",
            ),
        )

        with self.assertRaises(InvalidDocxError):
            process_document(document)

        document.refresh_from_db()

        self.assertEqual(
            document.status,
            Document.Status.FAILED,
        )

        self.assertEqual(
            document.text_content,
            "",
        )

        self.assertNotEqual(
            document.processing_error,
            "",
        )