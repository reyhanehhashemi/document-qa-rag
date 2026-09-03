from io import BytesIO

from django.test import SimpleTestCase
from docx import Document as WordDocument

from apps.documents.services.docx_parser import extract_docx_text
from apps.documents.services.exceptions import (
    EmptyDocxError,
    InvalidDocxError,
)


class DocxParserTests(SimpleTestCase):
    def create_docx_file(self):
        output = BytesIO()

        document = WordDocument()

        document.add_paragraph(
            "First paragraph"
        )

        table = document.add_table(
            rows=2,
            cols=2,
        )

        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "Alpha"
        table.cell(1, 1).text = "42"

        document.add_paragraph(
            "Last paragraph"
        )

        document.save(output)
        output.seek(0)

        return output

    def test_extracts_paragraphs_and_tables_in_order(self):
        source = self.create_docx_file()

        result = extract_docx_text(source)

        expected = (
            "First paragraph\n\n"
            "Name | Value\n"
            "Alpha | 42\n\n"
            "Last paragraph"
        )

        self.assertEqual(
            result,
            expected,
        )

    def test_rejects_invalid_docx_file(self):
        source = BytesIO(
            b"This is not a DOCX file."
        )

        with self.assertRaises(InvalidDocxError):
            extract_docx_text(source)

    def test_rejects_empty_docx_file(self):
        output = BytesIO()

        document = WordDocument()
        document.save(output)

        output.seek(0)

        with self.assertRaises(EmptyDocxError):
            extract_docx_text(output)