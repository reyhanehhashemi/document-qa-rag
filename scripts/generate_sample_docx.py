from pathlib import Path

from docx import Document


BASE_DIR = Path(
    __file__
).resolve().parent.parent

SAMPLE_DATA_DIR = (
    BASE_DIR / "sample_data"
)

OUTPUT_PATH = (
    SAMPLE_DATA_DIR
    / "northbridge_student_guide.docx"
)


def add_section(
    document,
    title,
    paragraphs,
):
    """
    Add one heading and its paragraphs to the DOCX document.
    """
    document.add_heading(
        title,
        level=1,
    )

    for paragraph in paragraphs:
        document.add_paragraph(
            paragraph
        )


def build_sample_document():
    """
    Generate the reproducible English DOCX used in demos
    and integration testing.
    """
    SAMPLE_DATA_DIR.mkdir(
        parents=True,
        exist_ok=True,
    )

    document = Document()

    document.add_heading(
        (
            "Northbridge University "
            "Student Services Guide"
        ),
        level=0,
    )

    document.add_paragraph(
        (
            "Test document for the "
            "Document QA RAG project."
        )
    )

    document.add_paragraph(
        (
            "This guide contains fictional university "
            "information created specifically for testing "
            "document upload, semantic retrieval, grounded "
            "question answering, source tracking, and "
            "question-answer history."
        )
    )

    add_section(
        document,
        "1. Course Registration",
        [
            (
                "Course registration opens fourteen days "
                "before the first day of each semester. "
                "Students must register through the Atlas "
                "Student Portal. Registration closes at "
                "5:00 PM on the third day of classes. "
                "Students may change their course selections "
                "until the same deadline."
            ),
            (
                "Advisor approval is not required for a "
                "normal study load. However, students who "
                "want to enroll in more than 18 credits must "
                "obtain written approval from their academic "
                "advisor before submitting the registration "
                "request."
            ),
        ],
    )

    add_section(
        document,
        "2. University Library",
        [
            (
                "The Central Library is open Monday through "
                "Friday from 8:00 AM to 8:00 PM. On Saturday, "
                "the library is open from 10:00 AM to "
                "4:00 PM. The library is closed on Sunday."
            ),
            (
                "Each student may borrow up to six books at "
                "one time. The standard loan period is "
                "fourteen days. A book may be renewed once "
                "if another student has not placed a hold "
                "on it."
            ),
        ],
    )

    add_section(
        document,
        "3. Tuition and Payments",
        [
            (
                "Semester tuition must be paid no later than "
                "five business days before the official start "
                "of the semester. Payments are made through "
                "the Finance section of the Atlas Student "
                "Portal."
            ),
            (
                "A late payment fee of 25 US dollars is "
                "applied when tuition is paid after the "
                "deadline. Students with an unpaid balance "
                "may be prevented from completing course "
                "registration."
            ),
        ],
    )

    add_section(
        document,
        "4. Student ID Cards",
        [
            (
                "New students can collect their university "
                "ID card from Room 204 in the Student "
                "Services Center after their enrollment has "
                "been verified. ID card collection is "
                "available Monday through Friday from "
                "9:00 AM to 3:00 PM."
            ),
            (
                "Students must bring a government-issued "
                "photo ID when collecting the university "
                "card. Replacement cards for lost IDs cost "
                "15 US dollars."
            ),
        ],
    )

    add_section(
        document,
        "5. IT Help Desk",
        [
            (
                "The IT Help Desk is located on the first "
                "floor of the Technology Building. In-person "
                "support is available Monday through Friday "
                "from 9:00 AM to 5:00 PM. Students can "
                "request password resets, Wi-Fi assistance, "
                "and access support for the Atlas Student "
                "Portal."
            ),
            (
                "The IT Help Desk does not provide support "
                "for privately owned printers or gaming "
                "devices."
            ),
        ],
    )

    add_section(
        document,
        "6. Important Scope Note",
        [
            (
                "This document does not contain information "
                "about cafeteria prices, campus housing fees, "
                "sports schedules, parking permit costs, or "
                "scholarship eligibility. Questions about "
                "those topics should be treated as unsupported "
                "by this document."
            ),
        ],
    )

    document.save(
        OUTPUT_PATH
    )

    return OUTPUT_PATH


if __name__ == "__main__":
    generated_path = (
        build_sample_document()
    )

    print(
        (
            "Sample DOCX generated at: "
            f"{generated_path}"
        )
    )